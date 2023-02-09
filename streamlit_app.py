import openai
import streamlit as st

# Inicializar la API de OpenAI
openai.api_key = "TU_API_KEY"

# Crear una función para hacer una llamada independiente a la API
def generate_text(prompt):
    completions = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=1000,
        n=1,
        stop=None,
        temperature=0.5,
    )

    message = completions.choices[0].text
    return message

# Crear la aplicación de Streamlit
st.title("Libro de buenos modales para ocasiones especiales")

# Tabla de contenido
st.header("Tabla de contenido")
st.write("1. Fiestas")
st.write("2. Bodas")
st.write("3. Funerales")
st.write("4. Comidas especiales")
st.write("5. Eventos formales")

# Capítulo 1: Fiestas
st.header("1. Fiestas")
prompt = "Escribe un capítulo sobre buenos modales en fiestas"
chapter_1 = generate_text(prompt)
st.write(chapter_1)

# Capítulo 2: Bodas
st.header("2. Bodas")
prompt = "Escribe un capítulo sobre buenos modales en bodas"
chapter_2 = generate_text(prompt)
st.write(chapter_2)

# Capítulo 3: Funerales
st.header("3. Funerales")
prompt = "Escribe un capítulo sobre buenos modales en funerales"
chapter_3 = generate_text(prompt)
st.write(chapter_3)

# Capítulo 4: Comidas especiales
st.header("4. Comidas especiales")
prompt = "Escribe un capítulo sobre buenos modales en comidas especiales"
chapter_4 = generate_text(prompt)
st.write(chapter_4)

# Capítulo 5: Eventos formales
st.header("5. Eventos formales")
prompt = "Escribe un capítulo sobre buenos modales en eventos formales"
chapter_5 = generate_text(prompt)
st.write(chapter_5)

Exportar el resultado a un documento
def export_to_docx():
# Instalar la librería python-docx si todavía no está instalada
!pip install python-docx
import docx

# Crear un nuevo documento de Word
doc = docx.Document()

# Agregar la tabla de contenido
doc.add_heading("Tabla de contenido", 0)
doc.add_paragraph("1. Fiestas", style='List Number')
doc.add_paragraph("2. Bodas", style='List Number')
doc.add_paragraph("3. Funerales", style='List Number')
doc.add_paragraph("4. Comidas especiales", style='List Number')
doc.add_paragraph("5. Eventos formales", style='List Number')

# Agregar los capítulos
doc.add_heading("1. Fiestas", 1)
doc.add_paragraph(chapter_1)
doc.add_heading("2. Bodas", 1)
doc.add_paragraph(chapter_2)
doc.add_heading("3. Funerales", 1)
doc.add_paragraph(chapter_3)
doc.add_heading("4. Comidas especiales", 1)
doc.add_paragraph(chapter_4)
doc.add_heading("5. Eventos formales", 1)
doc.add_paragraph(chapter_5)

# Guardar el documento
doc.save("libro_de_buenos_modales.docx")

# Agregar un botón para exportar el resultado a un documento
st.header("Exportar a un documento")
if st.button("Exportar a un documento"):
export_to_docx()
st.success("El documento se ha guardado como 'libro_de_buenos_modales.docx'")
